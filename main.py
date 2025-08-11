# main.py
import os
import io
import uuid
import asyncio
import logging
import requests
import PyPDF2
import numpy as np
from datetime import datetime
from typing import List, Optional, Any, Dict

from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document as DocxDocument

# LangChain imports (used for Document schema and text splitter)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.prompts import PromptTemplate

# sklearn for cosine similarity
from sklearn.metrics.pairwise import cosine_similarity

# Google Gen AI SDK
from google import genai
from google.genai import types

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

app = FastAPI(
    title="HackRx 6.0 - Enhanced LLM Query Retrieval System",
    description="Advanced document processing with LangChain and Vector Database for insurance, legal, HR, and compliance domains",
    version="2.0.0"
)

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

security = HTTPBearer()

# Environment-configurable defaults
DEFAULT_GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
DEFAULT_EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "gemini-embedding-001")


class QueryRequest(BaseModel):
    documents: str
    questions: List[str]
    use_advanced_chunking: Optional[bool] = True
    chunk_size: Optional[int] = 1000
    chunk_overlap: Optional[int] = 200
    retrieval_k: Optional[int] = 5


class QueryResponse(BaseModel):
    answers: List[str]
    document_id: str
    processing_time: float
    retrieval_method: str


class DocumentInfo(BaseModel):
    document_id: str
    document_url: str
    processed_at: str
    chunk_count: int
    embedding_method: str


def get_gemini_client() -> genai.Client:
    """Return a configured Gemini client (lazy init)."""
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY not found in environment variables")

    # Some SDK versions support genai.configure(); others accept api_key in Client()
    try:
        genai.configure(api_key=api_key)
        client = genai.Client()
    except Exception:
        client = genai.Client(api_key=api_key)
    return client


def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify API token from header against API_KEY env var."""
    api_key = os.getenv("API_KEY")
    if api_key is None or credentials.credentials != api_key:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid API key")
    return credentials.credentials


class AdvancedDocumentProcessor:
    """Enhanced document processor using LangChain-style splitting and Google GenAI client."""

    def __init__(self, client: Optional[genai.Client] = None):
        self.client = client or get_gemini_client()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2500,
            chunk_overlap=300,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )


    def extract_text_from_url(self, url: str) -> str:
        """Download and extract text from url/pdf/docx/text."""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            content_type = response.headers.get('content-type', '').lower()

            if url.lower().endswith('.pdf') or 'application/pdf' in content_type:
                return self._extract_text_from_pdf_bytes(response.content)
            elif url.lower().endswith('.docx') or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                return self._extract_text_from_docx_bytes(response.content)
            else:
                return response.text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error downloading document: {str(e)}")

    def _extract_text_from_pdf_bytes(self, data: bytes) -> str:
        pdf_file = io.BytesIO(data)
        reader = PyPDF2.PdfReader(pdf_file)
        text_pieces = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_pieces.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page {i}: {e}")
        return "\n".join(text_pieces)

    def _extract_text_from_docx_bytes(self, data: bytes) -> str:
        buf = io.BytesIO(data)
        doc = DocxDocument(buf)
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    def create_documents(self, text: str, source_url: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
        if not text or not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in the document")

        # update splitter params
        self.text_splitter.chunk_size = chunk_size
        self.text_splitter.chunk_overlap = chunk_overlap

        texts = self.text_splitter.split_text(text)
        documents: List[Document] = []
        for i, t in enumerate(texts):
            if len(t.strip()) > 50:
                documents.append(
                    Document(
                        page_content=t,
                        metadata={
                            "source": source_url,
                            "chunk_id": i,
                            "document_type": self._detect_document_type(source_url),
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                )
        return documents

    def _detect_document_type(self, url: str) -> str:
        u = url.lower()
        if '.pdf' in u:
            return 'pdf'
        if '.docx' in u:
            return 'docx'
        if '.txt' in u:
            return 'text'
        return 'unknown'


class VectorStoreManager:
    """In-memory vector store with batched + throttled Gemini embeddings."""

    def __init__(self, client: genai.Client):
        self.client = client
        self.stores: Dict[str, Dict[str, Any]] = {}

    async def create_vector_store(self, documents: List[Document], document_id: str, method: str = "in-memory") -> Dict:
        texts = [d.page_content for d in documents]
        if not texts:
            raise ValueError("No texts to embed")

        embeddings = []
        batch_size = 90  # Safe limit under Gemini's 100 max
        delay = 1.0      # seconds between batches
        model_name = os.getenv("EMBEDDING_MODEL", DEFAULT_EMBEDDING_MODEL)

        logger.info(f"Embedding {len(texts)} chunks for document {document_id} using {model_name}")

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i+batch_size]
            tries = 0
            while tries < 3:
                try:
                    res = self.client.models.batch_embed_contents(
                        model=model_name,
                        requests=[{"content": t} for t in batch]
                    )
                    batch_embeddings = [np.array(e.values) for e in res.embeddings]
                    embeddings.extend(batch_embeddings)
                    break  # success → break retry loop
                except Exception as e:
                    tries += 1
                    logger.error(f"Embedding batch {i//batch_size+1} failed (attempt {tries}): {e}")
                    if tries < 3:
                        await asyncio.sleep(2 ** tries)  # exponential backoff
                    else:
                        raise

            if i + batch_size < len(texts):
                await asyncio.sleep(delay)  # throttle to avoid 429s

        # Normalize
        normed = []
        for emb in embeddings:
            norm = np.linalg.norm(emb)
            normed.append(emb if norm == 0 else emb / norm)

        self.stores[document_id] = {
            "documents": documents,
            "embeddings": np.vstack(normed)
        }
        logger.info(f"Created vector store '{document_id}' with {len(documents)} chunks (batched & throttled)")
        return self.stores[document_id]



    def get_retriever(self, document_id: str, k: int = 5):
        store = self.stores.get(document_id)
        if not store:
            raise ValueError(f"No vector store found for document_id: {document_id}")
        return GenAICosineRetriever(store, self.client, k)


class GenAICosineRetriever:
    """Retriever using cosine similarity on stored embeddings."""

    def __init__(self, store: Dict[str, Any], client: genai.Client, k: int):
        self.documents: List[Document] = store["documents"]
        self.embeddings = store["embeddings"]
        self.client = client
        self.k = k

    def get_relevant_documents(self, query: str) -> List[Document]:
        if self.embeddings is None or self.embeddings.size == 0:
            logger.warning("Embeddings not available, falling back to text-based retrieval.")
            fr = FallbackRetriever({"documents": self.documents}, k=self.k)
            return fr.get_relevant_documents(query)

        # normal embedding retrieval
        res = self.client.models.embed_content(
            model=os.getenv("EMBEDDING_MODEL", DEFAULT_EMBEDDING_MODEL),
            contents=[query],
            config=types.EmbedContentConfig(task_type="SEMANTIC_SIMILARITY")
        )
        q_emb = np.array(res.embeddings[0].values)
        norm = np.linalg.norm(q_emb)
        if norm != 0:
            q_emb = q_emb / norm
        scores = (self.embeddings @ q_emb).tolist()
        top_idx = np.argsort(scores)[::-1][: self.k]
        return [self.documents[int(i)] for i in top_idx]

async def embed_in_batches(self, texts, model="models/text-embedding-004", batch_size=50, delay=1):
    all_embeddings = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        try:
            res = self.client.models.batch_embed_contents(
                model=model,
                requests=[{"content": t} for t in batch]
            )
            all_embeddings.extend([r.values for r in res.embeddings])
        except Exception as e:
            logger.error(f"Embedding batch failed: {e}")
        await asyncio.sleep(delay)  # avoid 429s
    return all_embeddings


class FallbackRetriever:
    """Simple word overlap fallback retriever."""

    def __init__(self, store: Dict[str, Any], k: int = 5):
        self.documents = store.get("documents", [])
        self.k = k

    def get_relevant_documents(self, query: str) -> List[Document]:
        q_words = set(query.lower().split())
        scored = []
        for d in self.documents:
            words = set(d.page_content.lower().split())
            overlap = len(q_words.intersection(words))
            if overlap > 0:
                score = overlap / max(1, len(q_words.union(words)))
                scored.append((score, d))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in scored[: self.k]]


class GeminiLLM:
    """Small wrapper around genai client for synchronous generation."""

    def __init__(self, client: genai.Client, model: str = DEFAULT_GEMINI_MODEL):
        self.client = client
        self.model = model

    def generate(self, prompt: str, max_output_tokens: int = 1024) -> str:
        try:
            resp = self.client.models.generate_content(
                model=os.getenv("GEMINI_MODEL", self.model),
                contents=prompt,
                # further params can be added via 'temperature', 'max_output_tokens', etc.
            )
            # Newer SDK: resp may have `.text`, `.candidates`, or `.output` structure.
            # We attempt to extract meaningful textual content robustly.
            if hasattr(resp, "text") and resp.text:
                return resp.text.strip()
            if hasattr(resp, "outputs") and resp.outputs:
                # outputs may be a list of generation objects with 'content' or 'text'
                out = resp.outputs[0]
                if hasattr(out, "content"):
                    return getattr(out, "content").strip()
                if isinstance(out, dict) and out.get("text"):
                    return out["text"].strip()
            if hasattr(resp, "candidates") and resp.candidates:
                return resp.candidates[0].content.strip()
            # As fallback stringify
            return str(resp).strip()
        except Exception as e:
            logger.error(f"Gemini generation failed: {e}")
            raise


class AdvancedQAChain:
    """Advanced QA chain using custom prompt + Gemini LLM wrapper."""

    def __init__(self, llm: Optional[GeminiLLM]):
        self.llm = llm
        self.qa_prompt = self._create_qa_prompt()

    def _create_qa_prompt(self) -> PromptTemplate:
        template = """
You are an expert analyst specializing in insurance, legal, HR, and compliance documents. 
Your task is to provide accurate, detailed, and contextually relevant answers based solely on the provided context.

Context Information:
{context}

Question: {question}

Instructions:
1. Base your answer ONLY on the information provided in the context above
2. If the information is not available in the context, respond with: "The requested information is not available in the provided document."
3. For legal and compliance topics, be precise and cite specific sections when possible
4. For insurance documents, focus on coverage details, exclusions, and policy terms
5. For HR documents, emphasize policies, procedures, and employee rights/responsibilities
6. Keep answers comprehensive yet concise (2-5 sentences)
7. Use professional terminology appropriate to the domain
8. If there are multiple relevant pieces of information, organize them clearly

Answer:
"""
        return PromptTemplate(template=template, input_variables=["context", "question"])

    def run(self, question: str, documents: List[Document]) -> str:
        if not documents or not question:
            return "The requested information is not available in the provided document."

        context = "\n\n".join([d.page_content for d in documents])
        formatted_prompt = self.qa_prompt.format(context=context, question=question)

        if not self.llm:
            # fallback simple extraction
            return self._fallback_answer(question, documents)

        try:
            answer = self.llm.generate(formatted_prompt)
            # The generation may contain extra whitespace; trim
            return answer.strip()
        except Exception as e:
            logger.error(f"QA generation failed: {e}")
            return self._fallback_answer(question, documents)

    def _fallback_answer(self, question: str, documents: List[Document]) -> str:
        # simple keyword based extraction
        q_words = set(question.lower().split())
        best_doc, best_score = None, 0
        for d in documents:
            words = set(d.page_content.lower().split())
            score = len(q_words.intersection(words))
            if score > best_score:
                best_score = score
                best_doc = d
        if not best_doc:
            return "The requested information is not available in the provided document."
        sentences = best_doc.page_content.split(". ")
        return ". ".join(sentences[:3]) + ('.' if len(sentences) > 3 else '')


# ---- Global instances ----
core_client = get_gemini_client()
doc_processor = AdvancedDocumentProcessor(client=core_client)
vector_manager = VectorStoreManager(core_client)
gemini_llm = GeminiLLM(core_client, model=os.getenv("GEMINI_MODEL", DEFAULT_GEMINI_MODEL))
qa_chain = AdvancedQAChain(gemini_llm)

# in-memory registry for processed documents
document_store: Dict[str, Dict[str, Any]] = {}


@app.post("/hackrx/run", response_model=QueryResponse)
async def process_query(request: QueryRequest, token: str = Depends(verify_token)):
    start_time = datetime.now()
    document_id = str(uuid.uuid4())

    try:
        logger.info(f"Processing document URL: {request.documents}")

        # Extract text
        document_text = doc_processor.extract_text_from_url(request.documents)

        # Chunk / create LangChain documents
        documents = doc_processor.create_documents(
            text=document_text,
            source_url=request.documents,
            chunk_size=request.chunk_size or 1000,
            chunk_overlap=request.chunk_overlap or 200
        )

        if not documents:
            raise HTTPException(status_code=400, detail="No valid document chunks created")

        # Create vector store (in-memory)
        try:
            vector_store = await vector_manager.create_vector_store(
                documents=documents,
                document_id=document_id,
                method="in-memory"
            )
        except Exception as e:
            logger.error(f"Embedding failed for {document_id}: {e}")
            vector_manager.stores[document_id] = {
                "documents": documents,
                "embeddings": None
            }
            retriever = FallbackRetriever(vector_manager.stores[document_id], k=request.retrieval_k or 5)
        else:
            retriever = vector_manager.get_retriever(document_id, k=request.retrieval_k or 5)


        # Save document info
        document_store[document_id] = {
            "url": request.documents,
            "processed_at": datetime.now().isoformat(),
            "chunk_count": len(documents),
            "embedding_method": "google-generative-ai"
        }

        # Retrieval + QA
        retriever = vector_manager.get_retriever(document_id, k=request.retrieval_k or 5)
        answers = []
        for q in request.questions:
            try:
                relevant_docs = retriever.get_relevant_documents(q)
                ans = qa_chain.run(q, relevant_docs)
                answers.append(ans)
            except Exception as e:
                logger.error(f"Error answering question '{q}': {e}")
                answers.append("Error processing this question. Please try again.")

        processing_time = (datetime.now() - start_time).total_seconds()
        return QueryResponse(
            answers=answers,
            document_id=document_id,
            processing_time=processing_time,
            retrieval_method="in-memory-cosine"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Error in process_query: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/document/{document_id}", response_model=DocumentInfo)
async def get_document_info(document_id: str, token: str = Depends(verify_token)):
    if document_id not in document_store:
        raise HTTPException(status_code=404, detail="Document not found")
    info = document_store[document_id]
    return DocumentInfo(
        document_id=document_id,
        document_url=info["url"],
        processed_at=info["processed_at"],
        chunk_count=info["chunk_count"],
        embedding_method=info.get("embedding_method", "unknown")
    )


@app.delete("/document/{document_id}")
async def delete_document(document_id: str, token: str = Depends(verify_token)):
    if document_id in document_store:
        del document_store[document_id]
    if document_id in vector_manager.stores:
        del vector_manager.stores[document_id]
    return {"message": f"Document {document_id} deleted successfully"}


@app.get("/documents")
async def list_documents(token: str = Depends(verify_token)):
    return {
        "documents": [
            {
                "document_id": doc_id,
                "url": info["url"],
                "processed_at": info["processed_at"],
                "chunk_count": info["chunk_count"]
            } for doc_id, info in document_store.items()
        ],
        "total_count": len(document_store)
    }


@app.get("/")
async def root():
    return {
        "message": "HackRx 6.0 - Enhanced LLM Query Retrieval System",
        "version": "2.0.0",
        "status": "running",
        "features": [
            "LangChain integration",
            "Vector database support (FAISS/Chroma optional)",
            "Advanced document chunking",
            "Domain-specific QA prompts",
            "Fallback mechanisms"
        ],
        "endpoint": "/hackrx/run"
    }


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "components": {
            "gemini_client": core_client is not None,
            "vector_stores_active": len(vector_manager.stores),
            "documents_processed": len(document_store)
        },
        "timestamp": datetime.now().isoformat()
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))
